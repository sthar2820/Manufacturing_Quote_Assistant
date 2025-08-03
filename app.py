"""
Manufacturing Quote Assistant - Streamlit Application
A comprehensive AI-powered tool for analyzing manufacturing quotes and documents.
"""

import os
import streamlit as st

# Load environment variables for local development
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass  # dotenv not available, skip loading .env file

# Core imports
try:
    import google.generativeai as genai
except ImportError as e:
    st.error(f"Failed to import Google Generative AI: {e}")
    st.error("Please ensure google-generativeai is installed: pip install google-generativeai")
    st.stop()

from datetime import datetime
from typing import Dict, Any, List, Optional
import pandas as pd
from pathlib import Path
import json
import tempfile
import zipfile

# Document processing imports
try:
    import fitz  # PyMuPDF
except ImportError:
    st.error("PyMuPDF not found. Please install: pip install PyMuPDF")
    st.stop()

try:
    from docx import Document
except ImportError:
    st.error("python-docx not found. Please install: pip install python-docx")
    st.stop()

try:
    import openpyxl
except ImportError:
    st.error("openpyxl not found. Please install: pip install openpyxl")
    st.stop()

# Document generation
from docx import Document as DocxDocument
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Vector store and embeddings (optional for RAG)
try:
    from langchain.embeddings import HuggingFaceEmbeddings
    from langchain.vectorstores import FAISS
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.chains import ConversationalRetrievalChain
    from langchain.llms.base import LLM
    from langchain.schema import BaseMessage
    RAG_AVAILABLE = True
except ImportError:
    RAG_AVAILABLE = False

# Configure Streamlit page
st.set_page_config(
    page_title="Manufacturing Quote Assistant",
    page_icon="🏭",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Constants
SUPPORTED_FILE_TYPES = ["pdf", "docx", "xlsx", "txt"]
OUTPUT_DIR = Path("processed_files")
OUTPUT_DIR.mkdir(exist_ok=True)

class ManufacturingQuoteAssistant:
    """Main application class for the Manufacturing Quote Assistant"""
    
    def __init__(self):
        """Initialize the application"""
        self.setup_api()
        if RAG_AVAILABLE:
            self.setup_rag()
        
    def setup_api(self):
        """Setup Google AI API"""
        # Try to get API key from Streamlit secrets first, then environment
        api_key = None
        try:
            api_key = st.secrets["GOOGLE_API_KEY"]
        except:
            api_key = os.getenv("GOOGLE_API_KEY")
        
        if not api_key:
            st.error("Google API Key not found. Please set GOOGLE_API_KEY in Streamlit secrets or environment.")
            st.info("For Streamlit Cloud: Add your API key in the app settings under Secrets")
            st.info("For local development: Add GOOGLE_API_KEY to your .env file")
            st.stop()
        
        try:
            genai.configure(api_key=api_key)
            self.model = genai.GenerativeModel('gemini-1.5-flash')
            st.success("AI Model initialized successfully")
        except Exception as e:
            st.error(f"Failed to initialize AI model: {e}")
            st.stop()
    
    def setup_rag(self):
        """Setup RAG components if available"""
        if not RAG_AVAILABLE:
            return
            
        try:
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2"
            )
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200
            )
            self.vector_store = None
            st.info("RAG capabilities enabled")
        except Exception as e:
            st.warning(f"RAG setup failed: {e}")
            RAG_AVAILABLE = False
    
    def extract_text_from_file(self, uploaded_file) -> str:
        """Extract text from uploaded file"""
        try:
            file_extension = uploaded_file.name.lower().split('.')[-1]
            
            if file_extension == "pdf":
                return self._extract_pdf_text(uploaded_file)
            elif file_extension == "docx":
                return self._extract_docx_text(uploaded_file)
            elif file_extension == "xlsx":
                return self._extract_xlsx_text(uploaded_file)
            elif file_extension == "txt":
                return str(uploaded_file.read(), "utf-8")
            else:
                st.error(f"Unsupported file type: {file_extension}")
                return ""
                
        except Exception as e:
            st.error(f"Error extracting text from {uploaded_file.name}: {e}")
            return ""
    
    def _extract_pdf_text(self, uploaded_file) -> str:
        """Extract text from PDF file"""
        text = ""
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
            tmp_file.write(uploaded_file.read())
            tmp_file.flush()
            
            doc = fitz.open(tmp_file.name)
            for page in doc:
                text += page.get_text()
            doc.close()
            
        os.unlink(tmp_file.name)
        return text
    
    def _extract_docx_text(self, uploaded_file) -> str:
        """Extract text from DOCX file"""
        doc = Document(uploaded_file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    
    def _extract_xlsx_text(self, uploaded_file) -> str:
        """Extract text from XLSX file"""
        workbook = openpyxl.load_workbook(uploaded_file)
        text = ""
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text += f"\n--- Sheet: {sheet_name} ---\n"
            
            for row in sheet.iter_rows():
                row_text = []
                for cell in row:
                    if cell.value is not None:
                        row_text.append(str(cell.value))
                if row_text:
                    text += " | ".join(row_text) + "\n"
        
        return text
    
    def analyze_quote_text(self, text: str, filename: str) -> Dict[str, Any]:
        """Analyze document text using AI"""
        prompt = self._build_analysis_prompt(text)
        
        try:
            with st.spinner(f"Analyzing {filename}..."):
                response = self.model.generate_content(
                    prompt,
                    generation_config=genai.types.GenerationConfig(
                        temperature=0.3,
                        max_output_tokens=4000,
                    )
                )
                
                if response.text:
                    return self._parse_analysis_output(response.text, filename)
                else:
                    st.warning("AI response was empty")
                    return self._create_empty_analysis(filename)
                    
        except Exception as e:
            st.error(f"AI analysis failed for {filename}: {e}")
            return self._create_empty_analysis(filename)
    
    def _build_analysis_prompt(self, text: str) -> str:
        """Build analysis prompt for AI"""
        return f"""
You are an expert manufacturing quote assistant focused on the Product Engineering Process.

Analyze this document to address these specific pain points:
1. Extract CRD (Customer Requirements Document) data
2. Identify requirements for comparison with past projects
3. Assess production feasibility concerns
4. Identify sourcing/purchasing requirements

Return your analysis using this EXACT structure:

---
CRD Data Summary
Extract and summarize: Part numbers, product specifications, quantities, delivery requirements, quality standards, and revision information.

Manufacturing Feasibility Assessment
Identify potential production challenges: Special processes (welding types, coating specs), unusual tolerances, equipment requirements, tooling needs, and any non-standard requirements.

Material & Component Sourcing Requirements
List: Raw material specifications (grades, dimensions, heat treatment), purchased components (fasteners, inserts, etc.), special materials, and supplier requirements.

Missing Critical Information
Identify what's missing for: Production planning, material sourcing, cost estimation, and feasibility assessment. Flag items that need clarification from customer or internal teams.

Comparison Baseline Data
Extract standardizable data points that could be compared with past projects: Material types, process requirements, tolerance ranges, quality specs, and production volumes.

Risk Factors & Special Requirements
Highlight: Non-standard processes, tight tolerances, special certifications, environmental requirements, and any factors that could impact production or sourcing.
---

Document content:
\"\"\"{text}\"\"\"
"""
    
    def _parse_analysis_output(self, output: str, filename: str) -> Dict[str, Any]:
        """Parse AI analysis output into structured data"""
        sections = {
            "CRD Data Summary": "",
            "Manufacturing Feasibility Assessment": "",
            "Material & Component Sourcing Requirements": "",
            "Missing Critical Information": "",
            "Comparison Baseline Data": "",
            "Risk Factors & Special Requirements": ""
        }
        
        # Simple section extraction
        current_section = None
        lines = output.split('\n')
        
        for line in lines:
            line = line.strip()
            
            # Check if this line is a section header
            for section_name in sections.keys():
                if section_name.lower() in line.lower():
                    current_section = section_name
                    break
            
            # Add content to current section
            if current_section and line and not any(s.lower() in line.lower() for s in sections.keys()):
                sections[current_section] += line + "\n"
        
        # Calculate risk assessment
        risk_data = self._assess_risk(sections["Missing Critical Information"])
        
        return {
            'file_name': filename,
            'crd_summary': sections["CRD Data Summary"].strip(),
            'feasibility_assessment': sections["Manufacturing Feasibility Assessment"].strip(),
            'sourcing_requirements': sections["Material & Component Sourcing Requirements"].strip(),
            'missing_info': sections["Missing Critical Information"].strip(),
            'baseline_data': sections["Comparison Baseline Data"].strip(),
            'risk_factors': sections["Risk Factors & Special Requirements"].strip(),
            'quote_ready': risk_data['quote_ready'],
            'risk_score': risk_data['risk_score'],
            'generated_at': datetime.now().isoformat(),
            'full_text': output
        }
    
    def _assess_risk(self, missing_info: str) -> Dict[str, Any]:
        """Assess risk based on missing information"""
        critical_missing = [
            "quantity", "material grade", "delivery", "bom", 
            "specifications", "tolerance", "dimensions", "process"
        ]
        
        missing_text_lower = missing_info.lower()
        missing_count = sum(1 for item in critical_missing if item in missing_text_lower)
        
        # Calculate risk score
        risk_score = min(20 + (missing_count * 15), 100)
        quote_ready = missing_count <= 1 and risk_score < 70
        
        return {
            'risk_score': risk_score,
            'quote_ready': quote_ready,
            'missing_count': missing_count
        }
    
    def _create_empty_analysis(self, filename: str) -> Dict[str, Any]:
        """Create empty analysis structure"""
        return {
            'file_name': filename,
            'crd_summary': "Analysis failed",
            'feasibility_assessment': "Analysis failed",
            'sourcing_requirements': "Analysis failed",
            'missing_info': "Analysis failed",
            'baseline_data': "Analysis failed",
            'risk_factors': "Analysis failed",
            'quote_ready': False,
            'risk_score': 100,
            'generated_at': datetime.now().isoformat(),
            'full_text': "Analysis failed"
        }
    
    def generate_word_summary(self, analyses: List[Dict[str, Any]]) -> str:
        """Generate Word document summary"""
        doc = DocxDocument()
        
        # Title
        title = doc.add_heading('Manufacturing Quote Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary information
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(f'Analysis Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Total Files Analyzed: {len(analyses)}')
        
        # Overall statistics
        quote_ready_count = sum(1 for analysis in analyses if analysis['quote_ready'])
        avg_risk_score = sum(analysis['risk_score'] for analysis in analyses) / len(analyses)
        
        doc.add_paragraph(f'Quote Ready Files: {quote_ready_count}/{len(analyses)}')
        doc.add_paragraph(f'Average Risk Score: {avg_risk_score:.1f}%')
        
        # Individual file analyses
        for i, analysis in enumerate(analyses, 1):
            doc.add_page_break()
            doc.add_heading(f'Analysis {i}: {analysis["file_name"]}', level=1)
            
            # Risk assessment
            risk_para = doc.add_paragraph()
            risk_para.add_run('Risk Score: ').bold = True
            risk_para.add_run(f'{analysis["risk_score"]}%')
            
            status_para = doc.add_paragraph()
            status_para.add_run('Quote Ready: ').bold = True
            status_para.add_run('Yes' if analysis['quote_ready'] else 'No')
            
            # Analysis sections
            sections = [
                ('CRD Data Summary', analysis['crd_summary']),
                ('Manufacturing Feasibility Assessment', analysis['feasibility_assessment']),
                ('Material & Component Sourcing Requirements', analysis['sourcing_requirements']),
                ('Missing Critical Information', analysis['missing_info']),
                ('Comparison Baseline Data', analysis['baseline_data']),
                ('Risk Factors & Special Requirements', analysis['risk_factors'])
            ]
            
            for section_title, content in sections:
                if content and content.strip():
                    doc.add_heading(section_title, level=2)
                    doc.add_paragraph(content)
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Quote_Analysis_Report_{timestamp}.docx"
        filepath = OUTPUT_DIR / filename
        doc.save(str(filepath))
        
        return str(filepath)
    
    def setup_rag_for_documents(self, texts: List[str]):
        """Setup RAG vector store for uploaded documents"""
        if not RAG_AVAILABLE:
            return False
            
        try:
            # Split texts into chunks
            all_chunks = []
            for text in texts:
                chunks = self.text_splitter.split_text(text)
                all_chunks.extend(chunks)
            
            # Create vector store
            self.vector_store = FAISS.from_texts(all_chunks, self.embeddings)
            return True
            
        except Exception as e:
            st.error(f"Failed to setup RAG: {e}")
            return False
    
    def chat_with_documents(self, question: str) -> str:
        """Chat with uploaded documents using RAG"""
        if not RAG_AVAILABLE or not self.vector_store:
            return "RAG not available. Please upload documents first."
        
        try:
            # Retrieve relevant documents
            docs = self.vector_store.similarity_search(question, k=3)
            context = "\n".join([doc.page_content for doc in docs])
            
            # Generate response using context
            prompt = f"""
Based on the following context from uploaded manufacturing documents, answer the question.

Context:
{context}

Question: {question}

Answer:
"""
            
            response = self.model.generate_content(prompt)
            return response.text if response.text else "Sorry, I couldn't generate a response."
            
        except Exception as e:
            return f"Error processing question: {e}"

def main():
    """Main application function"""
    st.title("Manufacturing Quote Assistant")
    st.markdown("### AI-Powered Document Analysis for Manufacturing Quotes")
    
    # Initialize the assistant
    if 'assistant' not in st.session_state:
        st.session_state.assistant = ManufacturingQuoteAssistant()
    
    assistant = st.session_state.assistant
    
    # Sidebar for file upload
    with st.sidebar:
        st.header("File Upload")
        uploaded_files = st.file_uploader(
            "Upload quote documents",
            type=SUPPORTED_FILE_TYPES,
            accept_multiple_files=True,
            help="Supported formats: PDF, DOCX, XLSX, TXT"
        )
        
        if uploaded_files:
            st.success(f"{len(uploaded_files)} file(s) uploaded")
            
            if st.button("Analyze Documents", type="primary"):
                st.rerun()
    
    # Main content area
    if uploaded_files:
        # Process documents
        with st.spinner("Processing documents..."):
            analyses = []
            extracted_texts = []
            
            for uploaded_file in uploaded_files:
                # Extract text
                text = assistant.extract_text_from_file(uploaded_file)
                if text:
                    extracted_texts.append(text)
                    # Analyze with AI
                    analysis = assistant.analyze_quote_text(text, uploaded_file.name)
                    analyses.append(analysis)
        
        if analyses:
            # Display results
            st.header("Analysis Results")
            
            # Summary metrics
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("Files Analyzed", len(analyses))
            
            with col2:
                quote_ready = sum(1 for a in analyses if a['quote_ready'])
                st.metric("Quote Ready", f"{quote_ready}/{len(analyses)}")
            
            with col3:
                avg_risk = sum(a['risk_score'] for a in analyses) / len(analyses)
                st.metric("Avg Risk Score", f"{avg_risk:.1f}%")
            
            with col4:
                if st.button("Generate Report"):
                    with st.spinner("Generating Word report..."):
                        report_path = assistant.generate_word_summary(analyses)
                        
                    # Provide download
                    with open(report_path, "rb") as file:
                        st.download_button(
                            label="Download Report",
                            data=file.read(),
                            file_name=os.path.basename(report_path),
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
            
            # Individual analysis results
            for i, analysis in enumerate(analyses):
                with st.expander(f"{analysis['file_name']} - Risk: {analysis['risk_score']}%"):
                    
                    # Status indicators
                    status_col1, status_col2 = st.columns(2)
                    with status_col1:
                        if analysis['quote_ready']:
                            st.success("Quote Ready")
                        else:
                            st.error("Not Quote Ready")
                    
                    with status_col2:
                        if analysis['risk_score'] < 30:
                            st.success(f"Low Risk ({analysis['risk_score']}%)")
                        elif analysis['risk_score'] < 70:
                            st.warning(f"Medium Risk ({analysis['risk_score']}%)")
                        else:
                            st.error(f"High Risk ({analysis['risk_score']}%)")
                    
                    # Analysis sections
                    sections = [
                        ("CRD Data Summary", analysis['crd_summary']),
                        ("Manufacturing Feasibility", analysis['feasibility_assessment']),
                        ("Sourcing Requirements", analysis['sourcing_requirements']),
                        ("Missing Information", analysis['missing_info']),
                        ("Baseline Data", analysis['baseline_data']),
                        ("Risk Factors", analysis['risk_factors'])
                    ]
                    
                    for section_title, content in sections:
                        if content and content.strip():
                            st.subheader(section_title)
                            st.write(content)
            
            # Setup RAG for chat
            if RAG_AVAILABLE and extracted_texts:
                if 'rag_setup' not in st.session_state:
                    with st.spinner("Setting up document chat..."):
                        rag_success = assistant.setup_rag_for_documents(extracted_texts)
                        st.session_state.rag_setup = rag_success
                
                if st.session_state.rag_setup:
                    # Chat interface
                    st.header("Chat with Documents")
                    st.markdown("Ask questions about your uploaded documents")
                    
                    if 'chat_history' not in st.session_state:
                        st.session_state.chat_history = []
                    
                    # Chat input
                    user_question = st.text_input("Ask a question about your documents:")
                    
                    if user_question:
                        with st.spinner("Thinking..."):
                            response = assistant.chat_with_documents(user_question)
                        
                        # Add to chat history
                        st.session_state.chat_history.append({
                            "question": user_question,
                            "answer": response,
                            "timestamp": datetime.now().strftime("%H:%M:%S")
                        })
                    
                    # Display chat history
                    for chat in reversed(st.session_state.chat_history[-5:]):  # Show last 5
                        st.markdown(f"**You ({chat['timestamp']}):** {chat['question']}")
                        st.markdown(f"**Assistant:** {chat['answer']}")
                        st.markdown("---")
    
    else:
        # Welcome message
        st.markdown("""
        ### Welcome to the Manufacturing Quote Assistant!
        
        This AI-powered tool helps you analyze manufacturing quotes and documents with:
        
        - **Multi-format Support**: PDF, DOCX, XLSX, TXT files
        - **AI Analysis**: Comprehensive quote evaluation using Google Gemini
        - **Risk Assessment**: Automated risk scoring and quote readiness
        - **Word Reports**: Professional analysis reports
        - **Document Chat**: Ask questions about your uploaded files
        
        **Get Started:**
        1. Upload your quote documents using the sidebar
        2. Click "Analyze Documents" to process them
        3. Review the analysis results and download reports
        4. Use the chat feature to ask specific questions
        """)

if __name__ == "__main__":
    main()
